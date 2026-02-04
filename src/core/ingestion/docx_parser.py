import logging
import re
from typing import List

from docx import Document as DocxDocument
from langchain_core.documents import Document

logger = logging.getLogger(__name__)


class DocxParser:
    """
    Парсер DOCX-документов, ориентированный на подготовку чанков
    для RAG-пайплайна (текст + таблицы с базовой семантикой).
    """

    def __init__(self, docx_path: str):
        # Загружаем DOCX-документ
        self.doc = DocxDocument(docx_path)

        # Текущий основной раздел документа (например: "1. ПРЕДМЕТ ДОГОВОРА")
        self.current_section = None

        # Регулярные выражения для удаления повторяющегося служебного текста
        # (артефакты экспорта, водяные знаки и т.п.)
        self.garbage_patterns = [
            r"\*\*ОГРАЖДАЮЩАЯ АКТУАЛЯЦИЯ\*\*",
            r"ОГРАЖДАЮЩАЯ АКТУАЛЯЦИЯ",
            r"\*\*ОГРАЖДАЮЩАЯ АКТУ\b",
        ]

    def parse(self) -> List[Document]:
        """
        Основная точка входа.

        Проходит по документу, нормализует текст,
        выделяет таблицы и смысловые блоки и
        возвращает список Document-объектов.
        """
        documents = []

        # Собираем весь текст документа в одну строку
        # (python-docx хранит текст по абзацам)
        full_text = ""
        for para in self.doc.paragraphs:
            full_text += para.text + "\n"

        # Нормализуем экранированные переводы строк
        full_text = full_text.replace("\\n", "\n")

        # Удаляем повторяющийся служебный / шумовой текст
        full_text = self._remove_garbage(full_text)

        # Разбиваем документ на строки для построчного анализа
        lines = full_text.split("\n")

        # Буфер для накопления обычного текстового контента
        current_content = []
        idx = 0

        while idx < len(lines):
            line = lines[idx].strip()

            # Пропускаем пустые строки
            if not line:
                idx += 1
                continue

            # Проверяем, начинается ли markdown-таблица
            if self._is_table_start(lines, idx):
                # Перед таблицей сохраняем накопленный текст как отдельный чанк
                if current_content:
                    text = "\n".join(current_content)
                    documents.append(self._create_document(text=text, chunk_type="paragraph"))
                    current_content = []

                # Извлекаем таблицу целиком
                table_md, new_idx = self._extract_table(lines, idx)
                documents.append(
                    self._create_document(text=table_md, chunk_type="table", is_atomic=True)
                )
                idx = new_idx
                continue

            # Проверяем, является ли строка заголовком основного раздела
            # (например: "1. ...", но не "1.1" или "2.3.4")
            if self._is_main_section(line):
                # Сохраняем предыдущий текстовый блок
                if current_content:
                    text = "\n".join(current_content)
                    documents.append(self._create_document(text=text, chunk_type="paragraph"))
                    current_content = []

                # Обновляем текущий раздел
                self.current_section = line

            # Добавляем строку в текущий текстовый блок
            current_content.append(line)
            idx += 1

        # Сохраняем оставшийся контент после завершения прохода
        if current_content:
            text = "\n".join(current_content)
            documents.append(self._create_document(text=text, chunk_type="paragraph"))

        return documents

    def _remove_garbage(self, text: str) -> str:
        """
        Удаляет повторяющийся служебный текст и
        нормализует количество пустых строк.
        """
        cleaned_text = text

        for pattern in self.garbage_patterns:
            cleaned_text = re.sub(pattern, "", cleaned_text, flags=re.IGNORECASE)

        # Схлопываем избыточные переводы строк
        cleaned_text = re.sub(r"\n{3,}", "\n\n", cleaned_text)

        return cleaned_text

    def _is_main_section(self, line: str) -> bool:
        """
        Определяет, является ли строка заголовком основного раздела.

        Поддерживаемые примеры:
          - "1. ПРЕДМЕТ ДОГОВОРА"
          - "10. РЕКВИЗИТЫ СТОРОН"
          - "Приложение №1"

        Исключаются вложенные подразделы:
          - "1.1. ..."
          - "2.3.4. ..."
        """
        if not line:
            return False

        # Нумерованный заголовок верхнего уровня
        if re.match(r"^(\d+)\.\s+([А-ЯЁ][А-ЯЁ\s]+)$", line):
            return True

        # Приложения
        if re.match(r"^(Приложение\s*№\s*\d+)", line, re.IGNORECASE):
            return True

        # Markdown-заголовки первого уровня
        if re.match(r"^#{1,2}\s+(.+)$", line):
            return True

        return False

    def _is_table_start(self, lines: List[str], idx: int) -> bool:
        """
        Проверяет, начинается ли с текущей строки markdown-таблица.
        """
        if idx >= len(lines) or idx + 1 >= len(lines):
            return False

        line1 = lines[idx].strip()
        line2 = lines[idx + 1].strip()

        # Таблица должна содержать разделители столбцов
        if "|" not in line1 or "|" not in line2:
            return False

        # Вторая строка — markdown-разделитель колонок
        if re.match(r"^\|[\s\-:|]+\|$", line2):
            return True

        return False

    def _extract_table(self, lines: List[str], start_idx: int):
        """
        Извлекает markdown-таблицу целиком.

        Возвращает:
          - текст таблицы
          - индекс строки, следующей за таблицей
        """
        table_lines = [lines[start_idx].strip()]
        idx = start_idx + 1

        while idx < len(lines):
            line = lines[idx].strip()
            if "|" in line:
                table_lines.append(line)
                idx += 1
            else:
                break

        return "\n".join(table_lines), idx

    def _create_document(self, text: str, chunk_type: str, is_atomic: bool = False) -> Document:
        """
        Формирует Document-объект с единым набором метаданных,
        используемых на этапе retrieval.
        """
        metadata = {
            "chunk_type": chunk_type,
            "section": self.current_section,
            "subsection": None,
        }

        # Атомарные чанки (например, таблицы) не должны дробиться дальше
        if is_atomic:
            metadata["is_atomic"] = True

        return Document(page_content=text, metadata=metadata)
